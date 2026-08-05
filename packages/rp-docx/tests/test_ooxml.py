"""The OOXML layer: namespaces, the package zip, and the `.dotx` content type."""

from __future__ import annotations

import zipfile

import docx
import pytest

from rp_docx import ooxml
from rp_docx.errors import InvalidDocxError, MissingFileError


class TestNamespaces:
    def test_qn_expands_a_known_prefix(self):
        assert ooxml.qn("w:t") == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"

    def test_qn_rejects_an_unknown_prefix(self):
        """A typo'd prefix is a bug in this package, not user input, so it must
        not silently produce an xpath that matches nothing."""
        with pytest.raises(KeyError, match="Unknown XML namespace prefix"):
            ooxml.qn("nope:t")

    def test_qn_passes_through_an_unprefixed_name(self):
        assert ooxml.qn("descr") == "descr"

    def test_xpath_binds_the_namespace_map(self, simple_docx):
        root = ooxml.parse_part(simple_docx, ooxml.DOCUMENT_PART)
        assert [node.text for node in ooxml.xpath(root, ".//w:t")][:2] == [
            "Title",
            "Alpha beta gamma.",
        ]


class TestReadability:
    def test_missing_file_is_an_input_error(self, tmp_path):
        with pytest.raises(MissingFileError):
            ooxml.check_readable(tmp_path / "nope.docx")

    def test_missing_file_is_also_a_filenotfounderror(self, tmp_path):
        """Library callers predating the suite hierarchy keep catching it."""
        with pytest.raises(FileNotFoundError):
            ooxml.check_readable(tmp_path / "nope.docx")

    def test_a_directory_is_an_input_error(self, tmp_path):
        with pytest.raises(MissingFileError):
            ooxml.check_readable(tmp_path)

    def test_non_zip_is_a_corrupt_file_error(self, not_a_docx):
        """python-docx conflates "missing" and "not a package"; we do not —
        they have different exit codes (1 and 3)."""
        with pytest.raises(InvalidDocxError):
            ooxml.check_readable(not_a_docx)

    def test_exit_codes_differ_for_the_two_failures(self, tmp_path, not_a_docx):
        with pytest.raises(MissingFileError) as missing:
            ooxml.check_readable(tmp_path / "nope.docx")
        with pytest.raises(InvalidDocxError) as corrupt:
            ooxml.check_readable(not_a_docx)
        assert missing.value.exit_code == 1
        assert corrupt.value.exit_code == 3


class TestParts:
    def test_read_part_returns_none_for_an_absent_part(self, simple_docx):
        """Optional parts are genuinely optional — commentsExtended.xml only
        exists once someone has resolved a comment."""
        assert ooxml.read_part(simple_docx, ooxml.COMMENTS_EXTENDED_PART) is None

    def test_part_names_lists_the_document(self, simple_docx):
        assert ooxml.DOCUMENT_PART in ooxml.part_names(simple_docx)

    def test_repack_substitutes_only_the_named_part(self, simple_docx, tmp_path):
        original = ooxml.read_part(simple_docx, ooxml.DOCUMENT_PART)
        target = ooxml.repack(simple_docx, tmp_path / "out.docx", {"docProps/app.xml": b"<a/>"})
        assert ooxml.read_part(target, "docProps/app.xml") == b"<a/>"
        assert ooxml.read_part(target, ooxml.DOCUMENT_PART) == original

    def test_repack_preserves_part_order(self, simple_docx, tmp_path):
        """Some OPC readers expect [Content_Types].xml first."""
        target = ooxml.repack(simple_docx, tmp_path / "out.docx", {})
        assert zipfile.ZipFile(target).namelist() == zipfile.ZipFile(simple_docx).namelist()

    def test_parse_part_rejects_malformed_xml(self, simple_docx, tmp_path):
        broken = ooxml.repack(simple_docx, tmp_path / "b.docx", {ooxml.DOCUMENT_PART: b"<w:p>"})
        with pytest.raises(InvalidDocxError, match="not well-formed"):
            ooxml.parse_part(broken, ooxml.DOCUMENT_PART)


class TestContentTypes:
    """Spec section 5.3, and the finding behind it.

    python-docx does not open a `.dotx` at all: it reads [Content_Types].xml,
    sees the template content type, and raises ValueError. Since house templates
    are the normal path, everything goes through ooxml.opened().
    """

    def test_a_saved_docx_declares_the_document_type(self, simple_docx):
        assert ooxml.content_type(simple_docx) == ooxml.DOCUMENT_CONTENT_TYPE
        assert not ooxml.is_template(simple_docx)

    def test_the_minimal_fixture_is_a_genuine_dotx(self, minimal_template):
        assert ooxml.content_type(minimal_template) == ooxml.TEMPLATE_CONTENT_TYPE
        assert ooxml.is_template(minimal_template)

    def test_python_docx_cannot_open_a_dotx_directly(self, minimal_template):
        """The finding, asserted rather than remembered. If a future python-docx
        learns to open templates, this fails and ooxml.opened() can be
        simplified — which is worth knowing about."""
        with pytest.raises(ValueError, match="not a Word file"):
            docx.Document(str(minimal_template))

    def test_opened_reads_a_dotx_anyway(self, minimal_template):
        with ooxml.opened(minimal_template) as document:
            assert "Normal" in {style.name for style in document.styles}

    def test_opened_reads_a_plain_docx(self, simple_docx):
        with ooxml.opened(simple_docx) as document:
            assert document.paragraphs[0].text == "Title"

    def test_opened_leaves_the_template_untouched(self, minimal_template):
        before = minimal_template.read_bytes()
        with ooxml.opened(minimal_template):
            pass
        assert minimal_template.read_bytes() == before

    def test_opened_maps_a_non_package_to_a_corrupt_file_error(self, not_a_docx):
        with pytest.raises(InvalidDocxError):
            with ooxml.opened(not_a_docx):
                pass

    def test_retyping_round_trips(self, simple_docx, tmp_path):
        as_template = ooxml.retype_as_template(simple_docx, tmp_path / "t.dotx")
        assert ooxml.is_template(as_template)
        back = ooxml.retype_as_document(as_template, tmp_path / "d.docx")
        assert not ooxml.is_template(back)

    def test_retyping_loses_no_parts(self, simple_docx, tmp_path):
        as_template = ooxml.retype_as_template(simple_docx, tmp_path / "t.dotx")
        assert zipfile.ZipFile(as_template).namelist() == zipfile.ZipFile(simple_docx).namelist()

    def test_retyping_in_place_rewrites_the_file(self, simple_docx, tmp_path):
        target = tmp_path / "copy.dotx"
        target.write_bytes(simple_docx.read_bytes())
        assert ooxml.retype_as_template(target) == target
        assert ooxml.is_template(target)

    def test_save_honours_a_dotx_extension(self, tmp_path):
        """python-docx always writes the document content type, so a file named
        .dotx would otherwise be a document wearing a template's extension."""
        written = ooxml.save(docx.Document(), tmp_path / "out.dotx")
        assert ooxml.is_template(written)

    def test_save_leaves_a_docx_as_a_document(self, tmp_path):
        written = ooxml.save(docx.Document(), tmp_path / "out.docx")
        assert not ooxml.is_template(written)
