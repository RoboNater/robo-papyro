"""Fixtures: every document and template used by the suite, built at run time.

**No binary templates in git** (spec section 11.1). A downloaded or corporate
`.dotx` in the repository is a licensing question, an opaque diff, and a
debugging hazard at once — when a test fails you cannot tell whether the code or
the template changed. Everything here is constructed from python-docx, so a
failure is always the code.

The three templates are **adversarial rather than realistic**, because realism
is not what catches bugs here:

``minimal``
    Word's built-in style names, Letter, no header. The happy path, and the
    path where the default ``StyleMap`` is correct.
``house_like``
    What a real house template looks like: non-Word style names, a name with a
    space and a non-ASCII character, A4, a header carrying an image, a linked
    character style, and a stylemap beside it.
``hostile``
    Missing a style the default ``StyleMap`` maps to, no stylemap, and two style
    names differing only by case. Exists to prove failures are loud.

The only committed binaries are the hand-made files under ``fixtures/`` that
python-docx genuinely cannot produce — tracked changes and comments (section
7). Those are built by editing XML directly rather than by python-docx's API,
which has none for either, so they too are generated here rather than committed.
"""

from __future__ import annotations

import functools
import io
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

import docx
import pytest
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Mm, Pt
from PIL import Image

from rp_docx import ooxml

#: Body text that must never appear in a manifest built from house_like — the
#: redaction assertion in spec section 5.2 searches for it.
SECRET_TEXT = "Zephyr Quokka Confidential Boilerplate"

HOUSE_STYLES = {
    "h1": "House Heading 1",
    "h2": "House Heading 2",
    "h3": "Résumé Heading",
    "h4": "House Heading 4",
    "body": "RP Body Text",
    "bullet": "House Bullet",
    "numbered": "House Number",
    "code": "House Code",
    "table": "Table Grid",
}

TABLE_DATA = [
    ["Region", "Units", "Revenue"],
    ["North", "120", "4,800"],
    ["South", "95", "3,610"],
]

IMAGE_SIZE = (48, 32)


@functools.lru_cache(maxsize=1)
def soffice_available() -> bool:
    """Whether LibreOffice is present **and can actually convert a document**.

    Checking only for the binary is not enough. A container can ship `soffice`
    that fails every conversion with "source file could not be loaded" — some
    minimal images lack the filters, others the Java runtime the filters expect.
    A presence check turns that into four confusing failures in an environment
    that never claimed to support conversion; a functional probe turns it into
    four honest skips, which `-rs` then reports.

    Probed once per session, on a document built for the purpose: nothing else
    tells the two cases apart.
    """
    if shutil.which("soffice") is None:
        return False
    with tempfile.TemporaryDirectory(prefix="rp-docx-probe-") as tmp:
        directory = Path(tmp)
        source = directory / "probe.docx"
        document = docx.Document()
        document.add_paragraph("probe")
        document.save(str(source))
        try:
            subprocess.run(
                [
                    "soffice",
                    f"-env:UserInstallation={(directory / 'profile').as_uri()}",
                    "--headless",
                    "--norestore",
                    "--invisible",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(directory),
                    str(source),
                ],
                capture_output=True,
                timeout=180,
            )
        except (OSError, subprocess.TimeoutExpired):
            return False
        return (directory / "probe.pdf").is_file()


def pytest_collection_modifyitems(config, items):
    """Skip ``@pytest.mark.requires_soffice`` tests when LibreOffice cannot work.

    A marker rather than an importable ``skipif``, because ``from conftest
    import ...`` only works when the tests directory happens to be on sys.path,
    which pytest's importlib import mode deliberately stops doing (see
    ci/test_workspace_invariants.py).
    """
    if not any("requires_soffice" in item.keywords for item in items):
        return
    if soffice_available():
        return
    skip = pytest.mark.skip(reason="LibreOffice (soffice) is absent or cannot convert documents")
    for item in items:
        if "requires_soffice" in item.keywords:
            item.add_marker(skip)


# --- helpers shared by tests, as fixtures rather than imports --------------


@pytest.fixture(scope="session")
def secret_text() -> str:
    return SECRET_TEXT


@pytest.fixture(scope="session")
def house_styles() -> dict[str, str]:
    return HOUSE_STYLES


@pytest.fixture(scope="session")
def table_data() -> list[list[str]]:
    return TABLE_DATA


@pytest.fixture(scope="session")
def run_cli():
    """Run the installed ``rp-docx`` console script and capture its output."""

    def run(*args) -> subprocess.CompletedProcess:
        return subprocess.run(
            ["rp-docx", *[str(a) for a in args]],
            capture_output=True,
            text=True,
            encoding="utf-8",
        )

    return run


@pytest.fixture(scope="session")
def run_umbrella():
    """Run ``rp docx ...`` — the same typer app, reached the other way."""

    def run(*args) -> subprocess.CompletedProcess:
        return subprocess.run(
            ["rp", "docx", *[str(a) for a in args]],
            capture_output=True,
            text=True,
            encoding="utf-8",
        )

    return run


@pytest.fixture
def cli_error():
    """Read the ErrorDetail out of a failed CLI run.

    The suite writes the human-readable message and then an ``rp_core``
    ErrorEnvelope to stderr, the envelope last so it survives any warnings the
    command printed first (spec section 4.1).
    """

    def parse(result) -> dict:
        return json.loads(result.stderr.splitlines()[-1])["error"]

    return parse


@pytest.fixture(scope="session")
def logo_png(tmp_path_factory) -> Path:
    path = tmp_path_factory.mktemp("assets") / "logo.png"
    Image.new("RGB", IMAGE_SIZE, (10, 120, 200)).save(path)
    return path


@pytest.fixture(scope="session")
def image_size() -> tuple[int, int]:
    return IMAGE_SIZE


# --- templates -------------------------------------------------------------


@pytest.fixture(scope="session")
def template_dir(tmp_path_factory) -> Path:
    return tmp_path_factory.mktemp("templates")


@pytest.fixture(scope="session")
def minimal_template(template_dir: Path) -> Path:
    """Built-in style names, Letter, no header. Saved as a genuine `.dotx`."""
    document = docx.Document()
    document.add_paragraph("Template body placeholder.")
    return ooxml.save(document, template_dir / "minimal.dotx")


def _add_paragraph_style(document, name: str, base: str | None = None):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base is not None:
        style.base_style = document.styles[base]
    return style


@pytest.fixture(scope="session")
def house_like_template(template_dir: Path, logo_png: Path) -> Path:
    """A stand-in for a real house template, with every awkward feature.

    Non-Word style names throughout, one carrying both a space and a non-ASCII
    character, A4 rather than Letter, a header with an image in it, a character
    style linked to a paragraph style, and a stylemap beside the file.
    """
    document = docx.Document()

    _add_paragraph_style(document, HOUSE_STYLES["h1"], "Heading 1")
    _add_paragraph_style(document, HOUSE_STYLES["h2"], "Heading 2")
    _add_paragraph_style(document, HOUSE_STYLES["h3"], "Heading 3")
    _add_paragraph_style(document, HOUSE_STYLES["h4"], "Heading 4")
    _add_paragraph_style(document, HOUSE_STYLES["body"], "Normal")
    _add_paragraph_style(document, HOUSE_STYLES["bullet"], "List Bullet")
    _add_paragraph_style(document, HOUSE_STYLES["numbered"], "List Number")
    _add_paragraph_style(document, HOUSE_STYLES["code"], "Normal")

    # A character style linked to a paragraph style — Word writes these in
    # pairs, and a reader that assumes every style is standalone trips on them.
    linked = document.styles.add_style("RP Body Text Char", WD_STYLE_TYPE.CHARACTER)
    linked.element.append(linked.element.makeelement(ooxml.qn("w:link"), {}))
    link = linked.element.find(ooxml.qn("w:link"))
    link.set(ooxml.qn("w:val"), "RPBodyText")

    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin = section.right_margin = Mm(25)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.add_run().add_picture(str(logo_png), width=Inches(1.2))
    header_paragraph.add_run(SECRET_TEXT)

    document.add_paragraph(SECRET_TEXT, style=HOUSE_STYLES["body"])

    path = ooxml.save(document, template_dir / "house_like.dotx")
    stylemap = {role: name for role, name in HOUSE_STYLES.items()}
    (template_dir / "house_like.stylemap.json").write_text(
        json.dumps(stylemap, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    return path


@pytest.fixture(scope="session")
def hostile_template(template_dir: Path) -> Path:
    """Missing "Heading 1", no stylemap, and two styles differing only by case.

    The default ``StyleMap`` maps ``h1`` to "Heading 1", so any markdown with a
    top-level heading must fail here — loudly, naming the style, listing what is
    available.
    """
    document = docx.Document()
    document.styles["Heading 1"].delete()
    _add_paragraph_style(document, "House Body")
    _add_paragraph_style(document, "house body")
    document.add_paragraph("Hostile template body.")
    return ooxml.save(document, template_dir / "hostile.dotx")


@pytest.fixture(scope="session")
def docx_twin_template(template_dir: Path) -> Path:
    """A `.docx` sharing ``minimal``'s bare name, to prove `.dotx` wins."""
    document = docx.Document()
    document.add_paragraph("The .docx twin, which resolution must not pick.")
    return ooxml.save(document, template_dir / "minimal.docx")


@pytest.fixture
def templates_env(monkeypatch, template_dir: Path):
    """Point template resolution at the fixture directory and nowhere else."""
    from rp_docx import templates

    monkeypatch.setenv(templates.TEMPLATE_DIR_ENV, str(template_dir))
    monkeypatch.delenv(templates.DEFAULT_TEMPLATE_ENV, raising=False)
    # repo_root() walks up from the cwd and would otherwise also find the
    # checkout's own templates/ directory, making the test depend on what
    # happens to be sitting in it.
    monkeypatch.setattr(templates, "repo_root", lambda start=None: None)
    return template_dir


# --- documents -------------------------------------------------------------


@pytest.fixture(scope="session")
def doc_dir(tmp_path_factory) -> Path:
    return tmp_path_factory.mktemp("documents")


@pytest.fixture(scope="session")
def rich_docx(doc_dir: Path, logo_png: Path) -> Path:
    """One document exercising most of the read surface.

    Headings at three levels, a styled multi-run paragraph, a list, two tables
    (one nested inside a cell), an inline image with alt text, a header and a
    footer, two sections, and core properties.
    """
    document = docx.Document()
    document.core_properties.title = "Quarterly Report"
    document.core_properties.author = "Test Author"
    document.core_properties.category = "reports"
    document.core_properties.keywords = "quarterly, revenue"

    document.add_paragraph("Quarterly Report", style="Heading 1")
    body = document.add_paragraph("Revenue rose ")
    bold = body.add_run("sharply")
    bold.bold = True
    bold.font.size = Pt(14)
    body.add_run(" this ")
    italic = body.add_run("quarter")
    italic.italic = True
    body.add_run(".")

    document.add_paragraph("Regional detail", style="Heading 2")
    document.add_paragraph("North led on units.", style="Normal")
    document.add_paragraph("First point", style="List Bullet")
    document.add_paragraph("Second point", style="List Bullet")
    document.add_paragraph("Notes", style="Heading 3")

    table = document.add_table(rows=len(TABLE_DATA), cols=3)
    table.style = "Table Grid"
    for row_index, row in enumerate(TABLE_DATA):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value

    nested_host = document.add_table(rows=1, cols=1)
    nested = nested_host.cell(0, 0).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "nested"
    nested.cell(0, 1).text = "cell"

    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(str(logo_png), width=Inches(0.8))

    section = document.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "Header text"
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Footer text"

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Appendix", style="Heading 2")

    path = doc_dir / "rich.docx"
    document.save(str(path))
    _set_alt_text(path, "Company logo")
    return path


def _set_alt_text(path: Path, description: str) -> None:
    """Give the document's first inline picture a description.

    python-docx has no alt-text API, and alt text is the sort of thing a reader
    must handle both with and without — so the fixture sets it directly.
    """
    root = ooxml.parse_part(path, ooxml.DOCUMENT_PART)
    for node in ooxml.xpath(root, ".//wp:docPr"):
        node.set("descr", description)
        break
    ooxml.repack(path, path.with_suffix(".tmp"), {ooxml.DOCUMENT_PART: _tostring(root)})
    path.with_suffix(".tmp").replace(path)


def _tostring(root) -> bytes:
    from lxml import etree

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


@pytest.fixture(scope="session")
def simple_docx(doc_dir: Path) -> Path:
    """Two headings and two paragraphs — enough for round-trip assertions."""
    document = docx.Document()
    document.add_paragraph("Title", style="Heading 1")
    document.add_paragraph("Alpha beta gamma.")
    document.add_paragraph("Section", style="Heading 2")
    document.add_paragraph("Delta epsilon.")
    path = doc_dir / "simple.docx"
    document.save(str(path))
    return path


@pytest.fixture(scope="session")
def split_runs_docx(doc_dir: Path) -> Path:
    """A document where every placeholder straddles run boundaries.

    Word splits a logical string across runs for reasons unrelated to meaning,
    so this is the shape that breaks a naive ``run.text.replace()`` — in the
    body, in a table cell, in the header, and in the footer (spec section 6).
    """
    document = docx.Document()

    paragraph = document.add_paragraph()
    for piece in ("Dear {{ ", "clie", "nt.na", "me }}, welcome."):
        paragraph.add_run(piece)

    formatted = document.add_paragraph()
    formatted.add_run("Total: {{ amo")
    tail = formatted.add_run("unt }} due")
    tail.bold = True

    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("Cell {{ cli")
    cell_paragraph.add_run("ent.name }}")

    section = document.sections[0]
    section.header.is_linked_to_previous = False
    header_paragraph = section.header.paragraphs[0]
    header_paragraph.add_run("Head {{ ci")
    header_paragraph.add_run("ty }}")
    section.footer.is_linked_to_previous = False
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.add_run("Foot {{ ci")
    footer_paragraph.add_run("ty }}")

    path = doc_dir / "split_runs.docx"
    document.save(str(path))
    return path


@pytest.fixture(scope="session")
def empty_docx(doc_dir: Path) -> Path:
    document = docx.Document()
    path = doc_dir / "empty.docx"
    document.save(str(path))
    return path


@pytest.fixture(scope="session")
def not_a_docx(doc_dir: Path) -> Path:
    path = doc_dir / "fake.docx"
    path.write_text("this is not a Word document", encoding="utf-8")
    return path


@pytest.fixture(scope="session")
def page_break_docx(doc_dir: Path) -> Path:
    """Three pages, for rendering and conversion tests."""
    document = docx.Document()
    for index in range(1, 4):
        paragraph = document.add_paragraph(f"Page {index}")
        if index < 3:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    path = doc_dir / "pages.docx"
    document.save(str(path))
    return path


# --- documents python-docx cannot make: comments and tracked changes -------

_TRACKED_BODY = """<w:p>
  <w:r><w:t xml:space="preserve">Kept text </w:t></w:r>
  <w:ins w:id="101" w:author="Ada Lovelace" w:date="2024-03-01T10:00:00Z">
    <w:r><w:t xml:space="preserve">inserted words </w:t></w:r>
  </w:ins>
  <w:del w:id="102" w:author="Grace Hopper" w:date="2024-03-02T11:30:00Z">
    <w:r><w:delText xml:space="preserve">removed words </w:delText></w:r>
  </w:del>
  <w:r><w:t>tail.</w:t></w:r>
</w:p>
<w:p>
  <w:pPr><w:rPr>
    <w:ins w:id="103" w:author="Ada Lovelace" w:date="2024-03-01T10:05:00Z"/>
  </w:rPr></w:pPr>
  <w:r><w:t>Second paragraph.</w:t></w:r>
</w:p>"""

_COMMENT_BODY = """<w:p>
  <w:commentRangeStart w:id="0"/>
  <w:r><w:t>Anchored sentence.</w:t></w:r>
  <w:commentRangeEnd w:id="0"/>
  <w:r><w:commentReference w:id="0"/></w:r>
</w:p>
<w:p w14:paraId="11111111">
  <w:commentRangeStart w:id="1"/>
  <w:r><w:t>Second anchor.</w:t></w:r>
  <w:commentRangeEnd w:id="1"/>
  <w:r><w:commentReference w:id="1"/></w:r>
</w:p>"""

_COMMENTS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{w}" xmlns:w14="{w14}">
  <w:comment w:id="0" w:author="Ada Lovelace" w:initials="AL"
             w:date="2024-03-01T09:00:00Z">
    <w:p w14:paraId="AAAAAAA1"><w:r><w:t>Please clarify this.</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="1" w:author="Grace Hopper" w:initials="GH"
             w:date="2024-03-02T09:30:00Z">
    <w:p w14:paraId="AAAAAAA2"><w:r><w:t>Resolved already.</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""

_COMMENTS_EXTENDED_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:commentsEx xmlns:w15="{w15}" xmlns:w="{w}">
  <w15:commentEx w15:paraId="AAAAAAA1" w15:done="0"/>
  <w15:commentEx w15:paraId="AAAAAAA2" w15:done="1"/>
</w15:commentsEx>
"""

_COMMENTS_RELATIONSHIP = (
    '<Relationship Id="rIdComments" '
    'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
    'Target="comments.xml"/>'
    '<Relationship Id="rIdCommentsEx" '
    'Type="http://schemas.microsoft.com/office/2011/relationships/commentsExtended" '
    'Target="commentsExtended.xml"/>'
)

_COMMENT_CONTENT_TYPES = (
    '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-'
    'officedocument.wordprocessingml.comments+xml"/>'
    '<Override PartName="/word/commentsExtended.xml" ContentType="application/vnd.ms-word.'
    'commentsExtended+xml"/>'
)


def _document_with_body(base: Path, target: Path, body_xml: str) -> Path:
    """Replace the document body with hand-written markup.

    python-docx has no API for tracked changes or comments (spec section 7), so
    the fixtures for them are written as XML. Generated rather than committed:
    a binary fixture in git is exactly the opaque diff section 11.1 rules out,
    and these are small enough to build honestly.
    """
    document = ooxml.parse_part(base, ooxml.DOCUMENT_PART)
    body = ooxml.xpath(document, "//w:body")[0]
    sect_pr = ooxml.xpath(body, "./w:sectPr")
    for child in list(body):
        body.remove(child)

    from lxml import etree

    wrapper = etree.fromstring(
        f'<w:wrap xmlns:w="{ooxml.NS["w"]}" xmlns:w14="{ooxml.NS["w14"]}">{body_xml}</w:wrap>'
    )
    for child in wrapper:
        body.append(child)
    for node in sect_pr:
        body.append(node)
    return ooxml.repack(base, target, {ooxml.DOCUMENT_PART: _tostring(document)})


@pytest.fixture(scope="session")
def tracked_changes_docx(doc_dir: Path) -> Path:
    """An insertion, a deletion, and a paragraph-mark format change.

    Two authors, so the ``--author`` filter has something to filter. Note the
    deletion carries ``w:delText`` rather than ``w:t``: a reader that looks only
    for ``w:t`` silently reports deletions as empty.
    """
    base = doc_dir / "tracked_base.docx"
    docx.Document().save(str(base))
    return _document_with_body(base, doc_dir / "tracked.docx", _TRACKED_BODY)


@pytest.fixture(scope="session")
def comments_docx(doc_dir: Path) -> Path:
    """Two comments, one resolved, with anchors and a commentsExtended part."""
    base = doc_dir / "comments_base.docx"
    docx.Document().save(str(base))
    staged = _document_with_body(base, doc_dir / "comments_staged.docx", _COMMENT_BODY)

    rels_name = "word/_rels/document.xml.rels"
    rels = ooxml.read_part(staged, rels_name).decode("utf-8")
    rels = rels.replace("</Relationships>", _COMMENTS_RELATIONSHIP + "</Relationships>")

    content_types = ooxml.read_part(staged, "[Content_Types].xml").decode("utf-8")
    content_types = content_types.replace("</Types>", _COMMENT_CONTENT_TYPES + "</Types>")

    target = doc_dir / "comments.docx"
    ooxml.repack(
        staged,
        target,
        {
            rels_name: rels.encode("utf-8"),
            "[Content_Types].xml": content_types.encode("utf-8"),
        },
    )
    # comments.xml and commentsExtended.xml are new parts, so they are appended
    # rather than substituted; repack() only replaces what already exists.
    import zipfile

    with zipfile.ZipFile(target, "a", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            ooxml.COMMENTS_PART, _COMMENTS_XML.format(w=ooxml.NS["w"], w14=ooxml.NS["w14"])
        )
        archive.writestr(
            ooxml.COMMENTS_EXTENDED_PART,
            _COMMENTS_EXTENDED_XML.format(w15=ooxml.NS["w15"], w=ooxml.NS["w"]),
        )
    return target


@pytest.fixture(scope="session")
def markdown_source() -> str:
    """Markdown covering everything the block parser must handle (spec section 9)."""
    return io.StringIO(
        "\n".join(
            [
                "# Report Title",
                "",
                "An intro paragraph with **bold**, *italic*, and `code` spans.",
                "",
                "## Findings",
                "",
                "- First bullet",
                "- Second bullet",
                "",
                "1. First step",
                "2. Second step",
                "",
                "### Detail",
                "",
                "| Region | Units |",
                "|---|---|",
                "| North | 120 |",
                "| South | 95 |",
                "",
                "---",
                "",
                "See [the docs](https://example.invalid/docs) for more.",
                "",
                "```",
                "print('hello')",
                "```",
                "",
                "#### Deep heading",
                "",
                "Final paragraph.",
            ]
        )
    ).getvalue()
