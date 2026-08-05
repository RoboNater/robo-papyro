"""House templates: finding them, describing them, and rebuilding them.

House templates are the normal path, not the exception (spec section 5).
:func:`resolve_template` is what every write entry point calls, and it always
returns a real file — falling back to python-docx's own bundled default rather
than to a ``None`` every caller would have to special-case.

**Manifests exist because the real templates cannot enter the repository.** A
:class:`~rp_docx.models.TemplateManifest` is a template's *shape* — style names,
page geometry, presence flags — and nothing else. :func:`build_manifest` runs
against the real file wherever it lives, the JSON is committed, and
:func:`synthesize` rebuilds a structurally equivalent template at test time. The
template itself never leaves the machine that holds it, and the manifest is
diffable, pasteable into an issue, and safe to share.

**Style resolution fails loudly.** :func:`require_style` raises rather than
substituting something plausible, because a document built with the wrong styles
looks right to everyone until it reaches review.

Nothing here prints, and nothing here imports typer.
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from rp_docx import ooxml
from rp_docx.errors import InvalidDocxError, TemplateError
from rp_docx.models import StyleDef, StyleMap, StyleType, TemplateInfo, TemplateManifest

#: Directory searched first for templates named by bare name. Suite convention:
#: package-specific settings carry the package's prefix (parent spec section 2).
TEMPLATE_DIR_ENV = "RP_DOCX_TEMPLATE_DIR"

#: Name (or path) of the template used when a caller asks for none.
DEFAULT_TEMPLATE_ENV = "RP_DOCX_TEMPLATE"

#: `.dotx` before `.docx`: a directory holding both means someone kept a working
#: copy beside the template, and the template is what was asked for.
TEMPLATE_SUFFIXES = (".dotx", ".docx")

#: python-docx spells OOXML's ``numbering`` style type ``LIST``. The models
#: report OOXML's name, because that is what is in the file.
_STYLE_TYPES: dict[int, StyleType] = {1: "paragraph", 2: "character", 3: "table", 4: "numbering"}

#: Page sizes recognized by name, keyed by (width, height) in twips. Anything
#: else is reported by its dimensions rather than guessed at.
_PAGE_SIZES: dict[tuple[int, int], str] = {
    (12240, 15840): "Letter",
    (15840, 12240): "Letter landscape",
    (12240, 20160): "Legal",
    (11906, 16838): "A4",
    (16838, 11906): "A4 landscape",
    (11907, 16839): "A4",
    (8419, 11906): "A5",
}

#: Manifest margin key → the python-docx section attribute holding it. The two
#: distances are not spelled ``*_margin`` there, and reading them as if they were
#: silently records a template's header position as absent rather than as wrong.
_MARGIN_KEYS: dict[str, str] = {
    "top": "top_margin",
    "bottom": "bottom_margin",
    "left": "left_margin",
    "right": "right_margin",
    "header": "header_distance",
    "footer": "footer_distance",
}


# --- resolution ------------------------------------------------------------


def builtin_template() -> Path:
    """python-docx's own bundled default template.

    Returned by :func:`resolve_template` when nothing is configured, so that
    every caller downstream has a real file to open and none of them needs a
    "no template" branch.
    """
    from docx.api import _default_docx_path

    return Path(_default_docx_path())


def repo_root(start: Path | None = None) -> Path | None:
    """The nearest ancestor of ``start`` that looks like the project checkout.

    Templates live in ``<repo>/templates/`` (spec section 5.1), which only means
    anything when running from a checkout — an installed wheel has no repo. The
    marker is a ``templates`` directory next to a ``.git`` or ``pyproject.toml``,
    so an unrelated ``templates/`` in some working directory is not mistaken for
    the suite's.
    """
    current = (start or Path.cwd()).resolve()
    for candidate in (current, *current.parents):
        if (candidate / "templates").is_dir() and (
            (candidate / ".git").exists() or (candidate / "pyproject.toml").is_file()
        ):
            return candidate
    return None


def template_dirs() -> list[Path]:
    """Every directory searched for a template named by bare name, in order.

    ``RP_DOCX_TEMPLATE_DIR`` first (and it may name several directories,
    separated the way ``PATH`` is), then the checkout's ``templates/local/`` and
    ``templates/``. ``local/`` comes first because it is the gitignored drop
    point for the *real* templates (spec section 11.1): when a name exists in
    both, the real one is the one meant.
    """
    dirs: list[Path] = []
    configured = os.environ.get(TEMPLATE_DIR_ENV, "")
    for entry in configured.split(os.pathsep):
        if entry.strip():
            dirs.append(Path(entry.strip()).expanduser())
    root = repo_root()
    if root is not None:
        dirs.extend([root / "templates" / "local", root / "templates"])
    return [d for d in dirs if d.is_dir()]


def _lookup(name: str) -> Path | None:
    for directory in template_dirs():
        for suffix in TEMPLATE_SUFFIXES:
            candidate = directory / f"{name}{suffix}"
            if candidate.is_file():
                return candidate
    return None


def available_template_names() -> list[str]:
    """Bare names resolvable by :func:`resolve_template`, de-duplicated."""
    names: list[str] = []
    for directory in template_dirs():
        for path in sorted(directory.iterdir()):
            if path.suffix.lower() in TEMPLATE_SUFFIXES and path.stem not in names:
                names.append(path.stem)
    return names


def resolve_template(name_or_path: str | Path | None = None) -> Path:
    """Find the template a caller means (spec section 5.1).

    1. An existing path is used as given
    2. A bare name resolves against :func:`template_dirs`, `.dotx` before `.docx`
    3. ``None`` uses ``RP_DOCX_TEMPLATE``, or python-docx's bundled default
    4. Anything else raises :class:`TemplateError` listing what *is* available
    """
    if name_or_path is None:
        configured = os.environ.get(DEFAULT_TEMPLATE_ENV, "").strip()
        return resolve_template(configured) if configured else builtin_template()

    as_path = Path(name_or_path)
    if as_path.is_file():
        return as_path

    # A path-shaped argument that does not exist is a wrong path, not a name to
    # look up: reporting "no template called ../drafts/memo.dotx" would send the
    # user hunting through the template directories for a typo in their own path.
    text = str(name_or_path)
    if as_path.suffix or os.sep in text or (os.altsep and os.altsep in text):
        raise TemplateError(f"No such template file: {text}")

    found = _lookup(text)
    if found is not None:
        return found

    available = available_template_names()
    known = ", ".join(available) if available else "none"
    searched = ", ".join(str(d) for d in template_dirs()) or "no template directories"
    raise TemplateError(
        f"No template called {text!r}. Available: {known}. "
        f"Searched: {searched}. Set {TEMPLATE_DIR_ENV} to add a directory, "
        "or pass a path to the template file."
    )


# --- inspection ------------------------------------------------------------


def _style_defs(document: Any) -> list[StyleDef]:
    defs: list[StyleDef] = []
    for style in document.styles:
        style_type = _STYLE_TYPES.get(int(style.type)) if style.type is not None else None
        if style_type is None:
            continue
        # python-docx's _NumberingStyle exposes neither base_style nor builtin —
        # numbering styles inherit from nothing in OOXML — so both are read
        # defensively rather than by type-switching on a private class.
        base = getattr(style, "base_style", None)
        defs.append(
            StyleDef(
                name=style.name,
                type=style_type,
                builtin=bool(getattr(style, "builtin", False)),
                base_style=base.name if base is not None else None,
            )
        )
    # Sorted so two inspections of the same template compare equal regardless of
    # the order styles.xml happens to list them in — which is what makes the
    # manifest round-trip assertion in spec section 11.2 meaningful.
    return sorted(defs, key=lambda s: (s.type, s.name))


def _page_size(section: Any) -> str:
    width, height = section.page_width, section.page_height
    if width is None or height is None:
        return "unknown"
    key = (width.twips, height.twips)
    named = _PAGE_SIZES.get(key)
    return named if named else f"{width.twips}x{height.twips} twips"


def _margins(section: Any) -> dict[str, int]:
    margins: dict[str, int] = {}
    for key, attribute in _MARGIN_KEYS.items():
        value = getattr(section, attribute, None)
        if value is not None:
            margins[key] = value.twips
    return margins


def _header_parts(path: Path) -> list[str]:
    return [n for n in ooxml.part_names(path) if n.startswith("word/header") and n.endswith(".xml")]


def _footer_parts(path: Path) -> list[str]:
    return [n for n in ooxml.part_names(path) if n.startswith("word/footer") and n.endswith(".xml")]


def _header_image_count(path: Path) -> int:
    """Images referenced by header parts.

    Counted from the XML rather than from python-docx, because merely *reading*
    ``section.header`` through python-docx creates the header part when it is
    absent — which would make inspecting a template modify it.
    """
    count = 0
    for part in _header_parts(path):
        root = ooxml.parse_part(path, part)
        if root is None:
            continue
        count += len(ooxml.xpath(root, ".//pic:pic | .//v:imagedata"))
    return count


def _has_letterhead(path: Path) -> bool:
    """Whether a header carries anything — an image or text.

    "Letterhead" is a judgement call and this is the defensible version of it: a
    template with an empty header has no letterhead, one with content does. The
    manifest reports ``header_image_count`` beside this so a caller that wants
    the stricter "has a logo" test can apply it.
    """
    for part in _header_parts(path):
        root = ooxml.parse_part(path, part)
        if root is None:
            continue
        if ooxml.xpath(root, ".//pic:pic | .//v:imagedata"):
            return True
        if any(node.text and node.text.strip() for node in ooxml.xpath(root, ".//w:t")):
            return True
    return False


def _default_paragraph_style(document: Any) -> str | None:
    """The style Word applies to a paragraph with no explicit one.

    Read from ``w:style/@w:default`` rather than from python-docx, which exposes
    no accessor for the flag. Usually "Normal", but a house template is exactly
    the kind of file that renames it.
    """
    for style in document.styles:
        if style.type is None or int(style.type) != 1:
            continue
        if ooxml.attr(style.element, "w:default") in ("1", "true", "on"):
            return style.name
    return None


def inspect_template(path: Path) -> TemplateInfo:
    """Read a template's styles and page setup without modifying it."""
    path = ooxml.check_readable(Path(path))
    with ooxml.opened(path) as document:
        styles = _style_defs(document)
        page_size = _page_size(document.sections[0]) if len(document.sections) else "unknown"
    return TemplateInfo(
        name=path.stem,
        path=path,
        format="dotx" if ooxml.is_template(path) else "docx",
        styles=styles,
        page_size=page_size,
        has_letterhead=_has_letterhead(path),
    )


def stylemap_path(template: Path) -> Path:
    """Where a template's stylemap sits: ``<stem>.stylemap.json`` beside it."""
    return Path(template).with_suffix(".stylemap.json")


def load_stylemap(template: Path) -> StyleMap:
    """The template's ``StyleMap``, or the built-in-names default.

    A malformed stylemap raises rather than falling back to the default: a
    stylemap exists precisely because the defaults are wrong for this template,
    so quietly using them would produce exactly the mis-styled document the
    file was added to prevent.
    """
    path = stylemap_path(template)
    if not path.is_file():
        return StyleMap()
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise TemplateError(f"{path.name} is not readable JSON: {exc}") from exc
    try:
        return StyleMap.model_validate(data)
    except Exception as exc:
        raise TemplateError(f"{path.name} is not a valid stylemap: {exc}") from exc


def style_names(document: Any, style_type: StyleType | None = None) -> list[str]:
    """Style names defined in ``document``, optionally of one type."""
    names = []
    for style in document.styles:
        if style.type is None:
            continue
        resolved = _STYLE_TYPES.get(int(style.type))
        if style_type is None or resolved == style_type:
            names.append(style.name)
    return sorted(names)


def require_style(
    document: Any, role: str, name: str, *, style_type: StyleType = "paragraph"
) -> str:
    """``name`` if the document defines it, else :class:`TemplateError`.

    **Never falls back.** Spec section 5.1: a silent substitution produces a
    document that looks wrong in ways nobody notices until review. The message
    names the missing style and lists what the template does have, because the
    fix is almost always a one-line correction to the stylemap.

    Checked at the point of use rather than up front. Word's built-in style set
    has no code style at all, so an eager check of the whole ``StyleMap`` would
    reject python-docx's own default template for a role most documents never
    use. A document that contains no code block does not need a code style.
    """
    available = style_names(document, style_type)
    if name in available:
        return name
    shown = ", ".join(available[:40])
    if len(available) > 40:
        shown += f", … and {len(available) - 40} more"
    raise TemplateError(
        f"The template has no {style_type} style called {name!r}, which is what "
        f"the stylemap maps {role!r} to. Available {style_type} styles: {shown}. "
        "Correct the template's .stylemap.json rather than relying on a fallback."
    )


def list_templates() -> list[TemplateInfo]:
    """Inspect every template resolvable by bare name.

    A template that cannot be read is skipped rather than failing the listing —
    one bad file in a shared template directory must not hide the rest.
    """
    infos: list[TemplateInfo] = []
    seen: set[str] = set()
    for directory in template_dirs():
        for path in sorted(directory.iterdir()):
            if path.suffix.lower() not in TEMPLATE_SUFFIXES or path.stem in seen:
                continue
            seen.add(path.stem)
            try:
                infos.append(inspect_template(path))
            except (InvalidDocxError, TemplateError):
                continue
    return infos


# --- manifests -------------------------------------------------------------


def build_manifest(path: Path) -> TemplateManifest:
    """Describe a template's shape, carrying none of its content.

    Redaction is a correctness property (spec section 5.2): nothing read here is
    document text, image bytes, an author name, or a path beyond the template's
    own basename. ``tests/test_templates.py`` asserts it against a template
    stuffed with distinctive text.
    """
    path = ooxml.check_readable(Path(path))
    with ooxml.opened(path) as document:
        sections = document.sections
        first = sections[0] if len(sections) else None
        manifest = TemplateManifest(
            name=path.stem,
            format="dotx" if ooxml.is_template(path) else "docx",
            styles=_style_defs(document),
            page_size=_page_size(first) if first is not None else "unknown",
            page_margins_twips=_margins(first) if first is not None else None,
            default_paragraph_style=_default_paragraph_style(document),
            has_letterhead=_has_letterhead(path),
            header_image_count=_header_image_count(path),
            footer_present=bool(_footer_parts(path)),
            section_count=len(sections),
        )
    stylemap = stylemap_path(path)
    if stylemap.is_file():
        manifest.stylemap = load_stylemap(path)
    return manifest


def load_manifest(path: Path) -> TemplateManifest:
    """Read a manifest from the JSON :func:`build_manifest` emits."""
    path = Path(path)
    if not path.is_file():
        raise TemplateError(f"No such manifest: {path}")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise TemplateError(f"{path.name} is not readable JSON: {exc}") from exc
    try:
        return TemplateManifest.model_validate(data)
    except Exception as exc:
        raise TemplateError(f"{path.name} is not a valid template manifest: {exc}") from exc


# --- synthesis -------------------------------------------------------------

_PLACEHOLDER_LOGO = (
    # 1x1 grey PNG. A placeholder header image only has to be an image: the
    # point is that the header part and its relationship exist, not what they
    # depict. Inline bytes rather than a fixture file, so synthesis works from
    # an installed wheel with no data files alongside it.
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00"
    b"\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xa8\xa9\xa9\x01\x00\x02\x9c\x01\x1c"
    b"\x14\xa5\x0f\xdf\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _apply_styles(document: Any, manifest: TemplateManifest) -> None:
    from docx.enum.style import WD_STYLE_TYPE

    type_enum = {
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "character": WD_STYLE_TYPE.CHARACTER,
        "table": WD_STYLE_TYPE.TABLE,
        "numbering": WD_STYLE_TYPE.LIST,
    }
    wanted = {(s.type, s.name): s for s in manifest.styles}

    for style in list(document.styles):
        if style.type is None:
            continue
        resolved = _STYLE_TYPES.get(int(style.type))
        if resolved is not None and (resolved, style.name) not in wanted:
            style.delete()

    present = {(_STYLE_TYPES.get(int(s.type)), s.name) for s in document.styles if s.type}
    for key, spec in wanted.items():
        if key in present:
            continue
        style = document.styles.add_style(spec.name, type_enum[spec.type])
        # add_style always marks its result custom. The manifest records what the
        # original said, and TemplateInfo.styles compares on it, so set it back.
        style.element.set(ooxml.qn("w:customStyle"), "0" if spec.builtin else "1")

    by_name = {s.name: s for s in document.styles}
    for spec in manifest.styles:
        if spec.base_style and spec.base_style in by_name and spec.name in by_name:
            try:
                by_name[spec.name].base_style = by_name[spec.base_style]
            except (ValueError, KeyError):
                # A base style of a different type is not expressible; the shape
                # of the style set is what synthesis is for, not its inheritance.
                continue


def _apply_page_setup(document: Any, manifest: TemplateManifest) -> None:
    from docx.shared import Twips

    dimensions = {name: size for size, name in _PAGE_SIZES.items()}
    size = dimensions.get(manifest.page_size)
    if size is None and "x" in manifest.page_size:
        try:
            width_text, _, rest = manifest.page_size.partition("x")
            size = (int(width_text), int(rest.split()[0]))
        except (ValueError, IndexError):
            size = None
    for section in document.sections:
        if size is not None:
            section.page_width, section.page_height = Twips(size[0]), Twips(size[1])
        for key, value in (manifest.page_margins_twips or {}).items():
            attribute = _MARGIN_KEYS.get(key)
            if attribute is not None:
                setattr(section, attribute, Twips(value))


def synthesize(manifest: TemplateManifest, output: Path) -> Path:
    """Rebuild a structurally equivalent template from a manifest.

    Reproduces style definitions, page size and margins, section count, and a
    placeholder header image when the manifest records a letterhead. It does
    **not** reproduce fonts, colors, or spacing: the goal is structural
    equivalence for testing style resolution, not visual fidelity (spec section
    5.2).
    """
    import docx
    from docx.enum.section import WD_SECTION
    from docx.shared import Inches

    output = Path(output)
    document = docx.Document()

    _apply_styles(document, manifest)
    for _ in range(max(0, manifest.section_count - len(document.sections))):
        document.add_section(WD_SECTION.NEW_PAGE)
    _apply_page_setup(document, manifest)

    if manifest.has_letterhead or manifest.header_image_count:
        import io

        header = document.sections[0].header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        for _ in range(max(1, manifest.header_image_count)):
            paragraph.add_run().add_picture(io.BytesIO(_PLACEHOLDER_LOGO), width=Inches(1))
    if manifest.footer_present:
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False

    written = ooxml.save(document, output)
    if manifest.stylemap is not None:
        stylemap_path(written).write_text(
            manifest.stylemap.model_dump_json(indent=2) + "\n", encoding="utf-8"
        )
    return written


# --- stylemap scaffolding --------------------------------------------------

#: Substrings that suggest a style plays a given role, most specific first.
#: Only ever used by :func:`scaffold_stylemap`, which is explicitly a guess.
_ROLE_HINTS: dict[str, tuple[str, ...]] = {
    "h1": ("heading 1", "head 1", "h1", "title"),
    "h2": ("heading 2", "head 2", "h2", "subtitle"),
    "h3": ("heading 3", "head 3", "h3"),
    "h4": ("heading 4", "head 4", "h4"),
    "body": ("body text", "body", "normal", "default"),
    "bullet": ("list bullet", "bullet", "list paragraph"),
    "numbered": ("list number", "numbered", "number"),
    "code": ("source code", "code", "monospace", "preformatted"),
    "table": ("table grid", "table"),
}


def scaffold_stylemap(path: Path) -> StyleMap:
    """A best-effort ``StyleMap`` for a template, for a human to correct.

    Matches style names against common patterns. **Never authoritative** — the
    CLI says so in its output, and spec section 10 requires it to. A generated
    stylemap that happens to be wrong is worse than none, because it looks
    reviewed.
    """
    info = inspect_template(path)
    paragraph_styles = [s.name for s in info.styles if s.type == "paragraph"]
    table_styles = [s.name for s in info.styles if s.type == "table"]
    scaffold = StyleMap()
    for role, hints in _ROLE_HINTS.items():
        pool = table_styles if role == "table" else paragraph_styles
        match = _best_match(pool, hints)
        if match is not None:
            setattr(scaffold, role, match)
    return scaffold


def _best_match(names: list[str], hints: tuple[str, ...]) -> str | None:
    lowered = [(name, name.lower()) for name in names]
    for hint in hints:
        exact = [name for name, low in lowered if low == hint]
        if exact:
            return exact[0]
        contains = [name for name, low in lowered if hint in low]
        if contains:
            return min(contains, key=len)
    return None


__all__ = [
    "DEFAULT_TEMPLATE_ENV",
    "TEMPLATE_DIR_ENV",
    "TEMPLATE_SUFFIXES",
    "available_template_names",
    "build_manifest",
    "builtin_template",
    "inspect_template",
    "list_templates",
    "load_manifest",
    "load_stylemap",
    "repo_root",
    "require_style",
    "resolve_template",
    "scaffold_stylemap",
    "style_names",
    "stylemap_path",
    "synthesize",
    "template_dirs",
]
